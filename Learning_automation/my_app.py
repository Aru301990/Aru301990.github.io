from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


#Create new document
Document=Document()


Phone_Number = '+917829990830'
Email = 'aru.vadde@gmail.com'


# Add heading
heading = Document.add_heading("", level=1)
run = heading.add_run("Aruna Vadde")
run.bold = True
run.font.size = Pt(28)  # Adjust size (20–32 works well for titles)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


Document.add_paragraph(Phone_Number + '|' + Email)

# Add summary
Document.add_heading('Professional_summary', level=2)
summary = """Results-oriented Software Test Engineer with 6+ years of experience in manual testing 
within Agile environments, specializing in Payroll, Engineering, and Manufacturing 
domains. Expertise in functional, regression, system, and UAT testing of web-based 
applications. Proficient in API testing using Postman and experienced with test 
management tools including Azure DevOps, JIRA, Allure TestOps, and TFS. Skilled in 
defect management, test planning, and cross-functional collaboration to ensure high-quality 
product delivery. Known for strong communication skills, adaptability, and a proactive 
approach to process improvement. """
Document.add_paragraph(summary)

# Add technical skills
Document.add_heading('Technical_skills', level=2)
skills = ['Manual_Testing', 'API_Testing', 'Test_Management_Tools', 'Defect_Tracking','Database_Testing', 'Agile_Methodologies']
Document.add_paragraph(', '.join(skills))

#add Professional experience
Document.add_heading('Professional_experience', level=2)
p=Document.add_paragraph()
company=p.add_run('Senior Software_Test_Engineer at Techversant infotech Pvt Ltd')
start_date=p.add_run (' (Oct 2021 - Present)\n')
p.add_run('Tools Used: Azure DevOps, Postman, JIRA, TFS\n').bold=True
p.add_run('Roles & Responsibilities:\n').bold=True

responsibilities = [
    """• Tested payroll modules for small businesses and partners, including company setup, 
employee/employer data, and payroll processing. 
• Validated tax calculations and verified complex scenarios such as statutory holiday pay, 
workers’ compensation, and vacation pay. 
• Designed and executed detailed test cases and test scenarios covering payroll workflows 
(onboarding, compensation, benefits, compliance). 
• Conducted API testing for payroll services using Postman and Swagger. 
• Reported, tracked, and validated defects in Azure DevOps and JIRA; collaborated with 
developers and cross-functional teams for quick resolution. 
• Performed end-to-end system testing and ensured accurate data flow across payroll, 
accounting, and reporting modules. 
• Participated in Agile ceremonies (sprint planning, daily stand-ups, retrospectives) and 
contributed QA insights during requirement grooming. 
• Conducted UAT support by preparing test data, guiding business users, and validating 
real-time payroll runs. 
• Ensured cross-browser compatibility and UI validation through exploratory and 
regression testing. 
• Mentored junior testers in defect reporting, test design, and QA best practices. 
• Validated builds in staging and production environments, ensuring release readiness. 
• Recommended process improvements in defect tracking and test execution workflows, 
contributing to overall QA efficiency. """]



# Add Education
Document.add_heading('Education', level=2)
edu=Document.add_paragraph()
edu.add_run('Master of Computer Science and Engineering').bold=True
edu.add_run(' Acharya Institute of Technology, Bangalore - Graduated July 2016')


# Save document
Document.save('CV.docx')
